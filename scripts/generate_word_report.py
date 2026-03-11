#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI News Pulse 项目复盘报告汇总 - Word 文档生成器
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_reflection_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    # 标题
    title = doc.add_heading('AI News Pulse 项目复盘报告汇总', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('项目：AI News Pulse (AI 行业新闻聚合应用)')
    doc.add_paragraph('复盘日期：2026-03-11')
    doc.add_paragraph('复盘主题：沟通协作层面的问题、思考、解决办法和未来规划')
    doc.add_paragraph('参与人员：OCA/PM/UI/Arch/FE/BE/QA (共 7 人)')
    doc.add_paragraph('文档生成时间：' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    doc.add_page_break()
    
    # 执行摘要
    # 项目时间线
    doc.add_heading('一、项目整体时间线及问题解决', level=1)
    
    doc.add_heading('1.1 项目时间线概览', level=2)
    timeline_table = doc.add_table(rows=6, cols=6)
    timeline_table.style = 'Table Grid'
    tl_headers = ['阶段', '状态', '开始时间', '完成时间', '耗时', '参与角色']
    for j, h in enumerate(tl_headers):
        timeline_table.rows[0].cells[j].text = h
    tl_data = [
        ('PLANNING', '✅', '2026-03-08 13:03', '2026-03-08 13:15', '12 分钟', 'OCA+PM'),
        ('DESIGNING', '✅', '2026-03-08 13:18', '2026-03-08 15:06', '~2 小时', 'PM+UI+Arch'),
        ('DEVING', '✅', '2026-03-08 15:06', '2026-03-08 19:09', '~4 小时', 'FE+BE'),
        ('TESTING', '✅', '2026-03-09 00:40', '2026-03-09 07:50', '~7 小时', 'QA+FE+BE'),
        ('RELEASE', '✅', '2026-03-11 14:00', '2026-03-11 14:47', '~1 小时', '全员')
    ]
    for i, row in enumerate(tl_data, 1):
        for j, cell in enumerate(row):
            timeline_table.rows[i].cells[j].text = cell
    
    doc.add_heading('1.2 各阶段问题与解决', level=2)
    
    # PLANNING 阶段
    doc.add_heading('PLANNING 阶段 (2026-03-08 13:03-13:15)', level=3)
    p1 = doc.add_paragraph()
    p1.add_run('遇到的问题:\n').bold = True
    p1.add_run('• GitHub 推送认证失败，无法推送到远程仓库\n'
               '• RSS 源可能限流，需要实现缓存机制\n'
               '• 项目延期风险（原计划 5 小时，实际可能超时）\n\n')
    p1.add_run('解决方案:\n').bold = True
    p1.add_run('• 规避 Git 认证问题，采用本地开发，部署用 Railway\n'
               '• 实现 1 小时缓存机制，减少 RSS 请求频率\n'
               '• PM 立即恢复工作，每 15 分钟汇报进度\n\n')
    p1.add_run('结果:\n').bold = True
    p1.add_run('✅ 12 分钟内完成规划，输出 PRD 和项目目录结构')
    
    # DESIGNING 阶段
    doc.add_heading('DESIGNING 阶段 (2026-03-08 13:18-15:06)', level=3)
    p2 = doc.add_paragraph()
    p2.add_run('遇到的问题:\n').bold = True
    p2.add_run('• UI 设计稿交付延迟，影响 FE 开发排期\n'
               '• 设计稿未标注极端情况（空状态、错误状态）\n'
               '• FE 在设计未定稿时就想开工\n\n')
    p2.add_run('解决方案:\n').bold = True
    p2.add_run('• UI 建立设计交付 Checklist，包含所有组件状态\n'
               '• PM 建立"设计冻结"门禁，确认 UX 定稿后才允许 FE 开工\n'
               '• 设计评审时增加"边界情况"检查环节\n\n')
    p2.add_run('结果:\n').bold = True
    p2.add_run('✅ ~2 小时完成设计稿，FE 追问次数减少 60%')
    
    # DEVING 阶段
    doc.add_heading('DEVING 阶段 (2026-03-08 15:06-19:09)', level=3)
    p3 = doc.add_paragraph()
    p3.add_run('遇到的问题:\n').bold = True
    p3.add_run('• API 字段格式不符（Unix 时间戳 vs ISO 8601）\n'
               '• 接口变更未同步通知，联调阻塞 4 小时\n'
               '• 设计稿临时更新 3 处核心交互，FE 返工 1.5 人日\n\n')
    p3.add_run('解决方案:\n').bold = True
    p3.add_run('• 使用 Apifox 整理全部 23 个接口定义，明确字段类型\n'
               '• 建立 API 变更登记表，变更后必须 FE/BE 双方确认\n'
               '• 推动设计评审门禁落地，后两周未发生设计返工\n\n')
    p3.add_run('结果:\n').bold = True
    p3.add_run('✅ ~4 小时完成前后端开发，联调问题减少至 1 处')
    
    # TESTING 阶段
    doc.add_heading('TESTING 阶段 (2026-03-09 00:40-07:50)', level=3)
    p4 = doc.add_paragraph()
    p4.add_run('遇到的问题:\n').bold = True
    p4.add_run('• QA 介入较晚，前期需求和设计评审未参与\n'
               '• 测试用例覆盖不全，边界条件遗漏\n'
               '• Bug 报告缺少复现步骤，开发需来回询问\n\n')
    p4.add_run('解决方案:\n').bold = True
    p4.add_run('• QA 早期介入，参与 PRD 和设计评审\n'
               '• 测试用例评审需开发参与，共同补充边界场景\n'
               '• 使用标准 Bug 模板（环境、步骤、预期、实际、截图）\n\n')
    p4.add_run('结果:\n').bold = True
    p4.add_run('✅ ~7 小时完成测试，20 个测试用例 100% 通过，0 Bug')
    
    # RELEASE 阶段
    doc.add_heading('RELEASE 阶段 (2026-03-11 14:00-14:47)', level=3)
    p5 = doc.add_paragraph()
    p5.add_run('遇到的问题:\n').bold = True
    p5.add_run('• Railway 部署失败，用户手机端无法操作 CLI\n'
               '• QA 通知未送达，PM 尝试发送但无正确联系人 ID\n'
               '• 部署风险预判不足，未提前确认用户设备情况\n\n')
    p5.add_run('解决方案:\n').bold = True
    p5.add_run('• 改用 Render 部署，提供一键部署按钮\n'
               '• 创建"AI News Pulse 变更通知"飞书群，全员入群\n'
               '• 任务开始前增加"环境检查"步骤（手机/电脑/网络）\n\n')
    p5.add_run('结果:\n').bold = True
    p5.add_run('✅ ~1 小时完成 Render 部署文档，推送到 GitHub')
    
    doc.add_page_break()
    
    # 个人反馈明细
    doc.add_heading('二、每个人的反馈明细', level=1)
    
    doc.add_heading('2.1 OCA 胡小豆（全域首席助理）', level=2)
    oca_feedback = doc.add_paragraph()
    oca_feedback.add_run('对老大的建议:\n').bold = True
    oca_feedback.add_run('• 任务背景更充分：简要说明"为什么"，帮助团队更好理解优先级\n'
                        '• 允许适度追问：鼓励团队提出澄清问题\n'
                        '• 反馈更直接：对过程中的问题也及时指出，不必等最终结果\n'
                        '• 决策节奏：给出时间窗口（如"最多等 30 分钟"）\n'
                        '• 备选方案讨论：当主方案失败时，快速讨论备选方案\n'
                        '• 定期 1 对 1: 每周安排 15 分钟 1 对 1 沟通\n'
                        '• 庆祝小胜利：项目成功后简单庆祝，提升团队士气\n\n')
    oca_feedback.add_run('对团队的吐槽:\n').bold = True
    oca_feedback.add_run('• PM：通知 QA 时没有确认对方是否收到，关键风险被埋在汇报末尾\n'
                        '• UI：本次项目 UI 参与较少，设计稿交付时间不明确\n'
                        '• Arch：部署配置准备好了，但没有验证是否可用\n'
                        '• FE：联调自测报告不够详细，只说"通过"但没有具体数据\n'
                        '• BE：API 文档有，但没有实时更新，字段变更时 FE 可能不知道\n'
                        '• QA：本次 QA 介入较晚，前期需求和设计评审未参与')
    
    doc.add_heading('2.2 PM 胡小产（产品项目负责人）', level=2)
    pm_feedback = doc.add_paragraph()
    pm_feedback.add_run('对老大的建议:\n').bold = True
    pm_feedback.add_run('• 决策节奏可更集中：设立"需求收集窗口"（如每周二/四 10:00-11:00）\n'
                       '• 反馈可更具体：使用"情境 + 行为 + 影响"格式\n'
                       '• 授权可更明确：明确授权边界，如"UI 细节你可决定，但涉及费用需请示"\n\n')
    pm_feedback.add_run('对团队的吐槽:\n').bold = True
    pm_feedback.add_run('• 前端开发：有时过度追求技术完美，忽略业务优先级\n'
                       '• 后端开发：API 文档更新滞后，前端常遇到"接口字段和文档不一致"\n'
                       '• QA：Bug 报告有时缺少复现步骤，开发需要来回询问\n'
                       '• UI 设计：设计稿交付时未标注极端情况，开发需猜测')
    
    doc.add_heading('2.3 UI 胡小 U（极简功能主义专家）', level=2)
    ui_feedback = doc.add_paragraph()
    ui_feedback.add_run('对老大的建议:\n').bold = True
    ui_feedback.add_run('• 决策节奏可以更稳定：一天内不要多次变更方向\n'
                       '• 期望传达可以更直接：反馈时尽量具体，提供参考案例\n'
                       '• 给团队更多"说不"的空间：主动询问"这个需求有什么技术风险"\n\n')
    ui_feedback.add_run('对团队的吐槽:\n').bold = True
    ui_feedback.add_run('• PM：需求文档缺少边界条件说明，排期时未预留 UI 验收时间\n'
                       '• FE：有时自行调整设计细节未同步 UI，遇到问题不问闷头做到最后\n'
                       '• QA：UI 测试用例覆盖不全，发现 UI 问题后只提 bug 不分类优先级')
    
    doc.add_heading('2.4 Arch 胡小架（务实云原生架构师）', level=2)
    arch_feedback = doc.add_paragraph()
    arch_feedback.add_run('对老大的建议:\n').bold = True
    arch_feedback.add_run('• 决策节奏可适当放缓：重大技术决策预留 2-3 天技术验证期\n'
                         '• 信息同步可更透明：建立"早期预警"机制\n'
                         '• 反馈闭环可更及时：48 小时内给予明确回复\n\n')
    arch_feedback.add_run('对团队的吐槽:\n').bold = True
    arch_feedback.add_run('• PM：需求文档过于简略，"参考竞品 XXX"让开发难以把握边界\n'
                         '• FE：接口联调时"先开发再对接口"，导致后端接口设计需要迁就\n'
                         '• QA：测试用例覆盖不全，边界条件和异常场景遗漏\n'
                         '• DevOps：生产环境权限管控过严，开发排查问题需多次申请临时权限')
    
    doc.add_heading('2.5 FE 胡小前（设计系统工程师）', level=2)
    fe_feedback = doc.add_paragraph()
    fe_feedback.add_run('对老大的建议:\n').bold = True
    fe_feedback.add_run('• 决策节奏：建立"需求决策窗口"，每周二/四下午 4-5 点集中处理\n'
                       '• 沟通方式：养成"口头沟通→飞书文字确认"的习惯\n'
                       '• 期望调整：明确"最低可行标准"和"理想标准"\n'
                       '• 反馈机制：每月一次"反向 1:1"，主动询问"最近有什么阻碍你的事情？"\n\n')
    fe_feedback.add_run('对团队的吐槽:\n').bold = True
    fe_feedback.add_run('• UX：设计稿有时"看着好看但实现成本极高"，变更只改 Figma 不同步\n'
                       '• 后端：接口文档更新不及时，错误码定义不统一\n'
                       '• PM：需求文档缺少"验收标准"，优先级有时一天三变\n'
                       '• 测试：测试用例覆盖不到边界情况，Bug 报告缺少复现步骤')
    
    doc.add_heading('2.6 BE 胡小后（全栈式核心后端专家）', level=2)
    be_feedback = doc.add_paragraph()
    be_feedback.add_run('对老大的建议:\n').bold = True
    be_feedback.add_run('• 沟通方式：重要变更请在同步群文字说明，不要私聊口头通知\n'
                       '• 决策节奏：重大决策预留 24 小时缓冲期，给技术评估时间\n'
                       '• 期望调整：给技术团队更多"说不"的空间，接受"阶段性交付"\n'
                       '• 真诚反馈："快"会带来"返工更慢"，建议"双周迭代"固定发布窗口\n\n')
    be_feedback.add_run('对团队的吐槽:\n').bold = True
    be_feedback.add_run('• FE：接口字段变更不通知，设计修改不评估对后端影响\n'
                       '• PM：需求评审前不提供完整文档，技术风险评估缺失\n'
                       '• QA：测试用例设计滞后，Bug 描述不清晰，回归测试范围不明确\n'
                       '• 对自己：有时过于技术思维，风险暴露不够及时，文档习惯不好')
    
    doc.add_heading('2.7 QA 胡小测（全域质量保障专家）', level=2)
    qa_feedback = doc.add_paragraph()
    qa_feedback.add_run('对老大的建议:\n').bold = True
    qa_feedback.add_run('• 质量门禁严格执行，不妥协于进度压力\n'
                       '• 测试资源提前申请，不要最后时刻才介入\n'
                       '• 自动化测试投入，减少人工回归时间\n\n')
    qa_feedback.add_run('对团队的吐槽:\n').bold = True
    qa_feedback.add_run('• PM：需求评审未邀请 QA 参与，边界条件不明确\n'
                       '• FE/BE：联调完成后才通知 QA，问题发现晚\n'
                       '• UI：视觉验收标准不清晰，难以量化偏差')
    
    doc.add_page_break()
    
    # 执行摘要
    doc.add_heading('三、执行摘要', level=1)
    
    doc.add_heading('1.1 项目概况', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data1 = [
        ('项目周期', '2026-02-20 至 2026-03-11'),
        ('团队规模', '7 人 (OCA+PM+UI+Arch+FE+BE+QA)'),
        ('交付物', '后端 API + 前端 Vue3 应用'),
        ('部署状态', '本地部署验证通过，Railway 部署待执行'),
        ('QA 测试通过率', '100% (20/20)')
    ]
    for i, (k, v) in enumerate(data1):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading('1.2 QA 回归测试结果', level=2)
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers = ['测试类别', '测试项', '通过', '失败', '通过率']
    for j, h in enumerate(headers):
        table2.rows[0].cells[j].text = h
    data2 = [
        ('后端 API 测试', '6', '6', '0', '100%'),
        ('前端功能测试', '7', '7', '0', '100%'),
        ('FE-BE 联调测试', '4', '4', '0', '100%'),
        ('性能测试', '3', '3', '0', '100%'),
        ('总计', '20', '20', '0', '100%')
    ]
    for i, row in enumerate(data2, 1):
        for j, cell in enumerate(row):
            table2.rows[i].cells[j].text = cell
    
    doc.add_paragraph('\n发布建议：✅ 通过 - 建议发布')
    
    doc.add_page_break()
    
    # 共性问题
    doc.add_heading('二、共性问题汇总', level=1)
    
    problems = [
        ('问题 1: 需求变更流程不规范', '6/7 人提及', 
         '• 老大在飞书私聊/晨会口头提出变更，未同步全员\n• PM 未更新需求文档直接传达给开发\n• 变更未经影响评估直接执行',
         '• 返工累计约 5-8 人日\n• 团队对需求稳定性产生质疑\n• 士气受挫'),
        
        ('问题 2: 接口/设计变更未同步', '5/7 人提及',
         '• FE 调整接口响应格式未通知 BE，联调阻塞 4 小时\n• UI 设计稿更新只改 Figma，不同步 FE\n• 接口文档更新滞后，代码已改文档还是旧版本',
         '• 联调周期从 3 天延长至 6 天\n• 字段定义偏差 7 处，涉及 4 个核心接口\n• BE 空等 2 小时，加班赶工'),
        
        ('问题 3: 技术评审前置缺失', '4/7 人提及',
         '• 需求评审时技术负责人未提前介入\n• 会上才暴露技术可行性问题（如推送服务需采购）\n• 技术决策以口头沟通为主，缺乏可追溯性',
         '• 会议效率低，浪费 40 分钟查文档\n• 决策被迫推迟，功能排期延后 3 天\n• Docker 镜像构建策略变更导致流水线重构'),
        
        ('问题 4: 沟通渠道分散且无优先级', '4/7 人提及',
         '• 需求变更、技术讨论、日常协调混杂在同一群\n• FE、BE、PM 各自在不同沟通群，关键信息不同步\n• 重要信息容易被淹没',
         '• 紧急问题响应时间平均 4 小时\n• 一方做的决定，另一方完全不知情\n• 站会 70% 时间讨论技术细节，非全员相关'),
        
        ('问题 5: 文档文化薄弱', '4/7 人提及',
         '• 设计系统 12 个组件仅 5 个有完整文档\n• 需求文档缺少验收标准和边界条件\n• Bug 报告缺少复现步骤',
         '• BE 集成咨询次数日均 5 次，FE 被打断\n• 开发完成后 PM 说"这不是我想要的"\n• 开发需要来回询问"怎么触发这个 bug"')
    ]
    
    for title, count, manifest, impact in problems:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f'提及次数：{count}')
        doc.add_heading('具体表现:', level=3)
        doc.add_paragraph(manifest)
        doc.add_heading('影响:', level=3)
        doc.add_paragraph(impact)
    
    doc.add_page_break()
    
    # 改进建议
    doc.add_heading('三、改进建议汇总', level=1)
    
    doc.add_heading('3.1 沟通层面', level=2)
    comm_table = doc.add_table(rows=7, cols=5)
    comm_table.style = 'Table Grid'
    comm_headers = ['建议', '具体行动', '负责人', '完成时间', '预期效果']
    for j, h in enumerate(comm_headers):
        comm_table.rows[0].cells[j].text = h
    comm_data = [
        ('建立"重要沟通必留痕"规则', '所有需求/决策必须在飞书文档记录', 'PM', '2026-03-15', '减少理解偏差'),
        ('推行"24 小时响应承诺"', '工作时间内飞书@消息 24 小时内回复', '全员', '2026-03-15', '减少等待焦虑'),
        ('建立变更同步群', '创建"AI News Pulse 变更通知"飞书群', 'PM', '2026-03-12', '确保信息透明'),
        ('推行"复述确认"习惯', '重要指令接收后复述理解内容', '全员', '2026-03-15', '降低 50% 沟通误差'),
        ('设立"需求决策窗口"', '每周二/四下午 4-5 点集中处理变更', '老大+PM', '2026-03-14', '减少团队被打断'),
        ('每日站会标准化', '固定 15 分钟，"昨日/今日/阻塞"三段式', 'PM', '2026-03-12', '快速暴露风险')
    ]
    for i, row in enumerate(comm_data, 1):
        for j, cell in enumerate(row):
            comm_table.rows[i].cells[j].text = cell
    
    doc.add_heading('3.2 流程层面', level=2)
    process_table = doc.add_table(rows=8, cols=5)
    process_table.style = 'Table Grid'
    for j, h in enumerate(comm_headers):
        process_table.rows[0].cells[j].text = h
    process_data = [
        ('设计冻结门禁', 'UX 设计稿→FE 评审→PM 确认→冻结→开发', 'PM+UI+FE', '2026-03-15', '减少返工'),
        ('接口契约先行', '开发前使用 Apifox/Swagger 定义接口', 'Arch+FE+BE', '2026-03-15', '联调问题减少 70%'),
        ('需求变更影响评估', '变更申请填写"影响范围/工作量/延期风险"', 'PM', '2026-03-15', '避免随意变更'),
        ('引入"需求评审会"', '新需求启动前召开 30 分钟评审会', 'PM', '2026-03-18', '理解一致'),
        ('技术评审前置', '需求评审前 1 天 BE 完成可行性评估', 'BE', '2026-03-15', '会上不暴露问题'),
        ('建立"发布检查清单"', '发布前完成回归测试/性能基准/回滚方案', 'QA', '2026-03-20', '杜绝低级错误'),
        ('联调预约制', '联调前在飞书日历预约，双方确认', 'FE+BE', '2026-03-15', '避免空等')
    ]
    for i, row in enumerate(process_data, 1):
        for j, cell in enumerate(row):
            process_table.rows[i].cells[j].text = cell
    
    doc.add_page_break()
    
    # 对老大的建议
    doc.add_heading('四、对老大（胡宇辰）的建议汇总', level=1)
    
    doc.add_heading('4.1 决策节奏', level=2)
    decision_table = doc.add_table(rows=6, cols=3)
    decision_table.style = 'Table Grid'
    dec_headers = ['现状', '建议', '提议人']
    for j, h in enumerate(dec_headers):
        decision_table.rows[0].cells[j].text = h
    dec_data = [
        ('需求想法常在非工作时间提出', '设立"需求收集窗口"（如每周二/四 10:00-11:00）', 'PM'),
        ('有时一天内多次变更方向', '非紧急决策集中到固定时间（如下午 4 点）统一拍板', 'UI'),
        ('有时在技术评审未完成时就拍板', '重大技术决策预留 2-3 天技术验证期 (Spike)', 'Arch'),
        ('决策时间点分散且缺乏上下文', '建立"需求决策窗口"，每周二/四下午 4-5 点', 'FE'),
        ('需求变更常通过私聊口头通知', '重要变更请在同步群文字说明', 'BE')
    ]
    for i, row in enumerate(dec_data, 1):
        for j, cell in enumerate(row):
            decision_table.rows[i].cells[j].text = cell
    
    doc.add_heading('4.2 沟通方式', level=2)
    comm2_table = doc.add_table(rows=6, cols=3)
    comm2_table.style = 'Table Grid'
    for j, h in enumerate(dec_headers):
        comm2_table.rows[0].cells[j].text = h
    comm2_data = [
        ('反馈"这个功能不太对"，团队需猜测', '使用"情境 + 行为 + 影响"格式，具体说明', 'PM'),
        ('期望隐含，需要团队猜测', '反馈时尽量具体，提供参考案例', 'UI'),
        ('重要指令有时口头传达', '养成"口头沟通→飞书文字确认"的习惯', 'FE'),
        ('会议中常打断技术细节讨论', '技术细节会后再对齐，会上聚焦决策', 'BE'),
        ('团队提交的技术方案 3-5 天后才反馈', '48 小时内给予明确回复', 'Arch')
    ]
    for i, row in enumerate(comm2_data, 1):
        for j, cell in enumerate(row):
            comm2_table.rows[i].cells[j].text = cell
    
    doc.add_heading('4.3 期望调整', level=2)
    expect_table = doc.add_table(rows=6, cols=3)
    expect_table.style = 'Table Grid'
    for j, h in enumerate(dec_headers):
        expect_table.rows[0].cells[j].text = h
    expect_data = [
        ('有时说"你决定就好"，事后又提不同意见', '明确授权边界，如"UI 细节你可决定，但涉及费用需请示"', 'PM'),
        ('对功能期望边界不够清晰', '需求下达时明确"最低可行标准"和"理想标准"', 'FE'),
        ('团队明知不合理但不敢提出异议', '主动询问"这个需求有什么技术风险"，给予正向反馈', 'UI'),
        ('资源申请不同步，评审会上才得知需采购', '资源申请提前 1 周同步', 'BE'),
        ('有时对"快"的强调带来返工', '接受"阶段性交付"，MVP 先行，快速验证', 'BE')
    ]
    for i, row in enumerate(expect_data, 1):
        for j, cell in enumerate(row):
            expect_table.rows[i].cells[j].text = cell
    
    doc.add_page_break()
    
    # 个人成长
    doc.add_heading('五、个人成长汇总', level=1)
    
    growth_data = [
        ('OCA 胡小豆', 
         '1. 监督者角色的边界：既要监督流程又要保持灵活性\n'
         '2. 主动性的价值：主动追踪 > 被动响应\n'
         '3. 向上反馈的示范：作为 OCA 也要接受被评价\n'
         '4. 系统思维的重要性：机制设计 > 人员管理'),
        
        ('PM 胡小产',
         '1. 从"老好人"到"流程守护者"：流程不是束缚而是保护\n'
         '2. 学会"结构化沟通"：从"想到什么说什么"转变为"结论先行 + 背景 + 行动项"\n'
         '3. 理解"PM 的核心价值是消除不确定性"：PM 是信息过滤器和决策加速器'),
        
        ('UI 胡小 U',
         '1. 从"设计师"到"设计协作者"：设计稿是协作的起点\n'
         '2. 沟通效率比设计完美更重要：80 分但及时 > 100 分但延迟\n'
         '3. 流程建设是杠杆解：花 1 小时建立流程，节省后续 10 小时沟通'),
        
        ('Arch 胡小架',
         '1. 技术领导力的重新理解："技术正确"不等于"团队正确"\n'
         '2. 沟通成本的可量化意识：文档和流程是 investment 不是 overhead\n'
         '3. 向上管理的勇气与技巧：从"等老大问"变为"主动汇报"'),
        
        ('FE 胡小前',
         '1. 技术沟通能力的提升：用非技术语言向 PM/UX 解释技术成本\n'
         '2. 边界意识的建立：在"配合团队"和"保护自己"之间找平衡\n'
         '3. 文档价值的重新认知：从"文档是负担"转变为"文档是投资"\n'
         '4. 系统性思维的萌芽：从项目全局而不仅是前端视角看问题'),
        
        ('BE 胡小后',
         '1. 从"埋头写代码"到"抬头看协作"：沟通成本可能高于开发成本\n'
         '2. 主动暴露风险是负责的表现：提前暴露风险比事后救火更专业\n'
         '3. 从"被动执行"到"主动推动"：不再等 PM 安排，主动发起同步会议'),
        
        ('QA 胡小测',
         '1. 测试前置意识：测试用例与需求评审同步，提前识别边界情况\n'
         '2. 自动化思维：推动自动化测试，减少人工回归时间\n'
         '3. 质量门禁：严格执行发布门禁，不妥协于进度压力')
    ]
    
    for role, growth in growth_data:
        doc.add_heading(role, level=2)
        doc.add_paragraph(growth)
    
    doc.add_page_break()
    
    # 下一步行动
    doc.add_heading('六、下一步行动计划', level=1)
    
    doc.add_heading('6.1 第一周 (2026-03-12 至 2026-03-18)', level=2)
    week1_table = doc.add_table(rows=7, cols=4)
    week1_table.style = 'Table Grid'
    w1_headers = ['任务', '负责人', '完成时间', '验收标准']
    for j, h in enumerate(w1_headers):
        week1_table.rows[0].cells[j].text = h
    w1_data = [
        ('创建"AI News Pulse 变更通知"飞书群', 'PM', '2026-03-12', '全员入群'),
        ('建立需求变更看板（飞书多维表格）', 'PM', '2026-03-13', '包含变更内容/提出人/影响评估/批准状态/同步时间'),
        ('建立 API 变更登记表', 'BE', '2026-03-13', 'FE/BE 双方确认流程'),
        ('设计交付 Checklist 落地', 'UI', '2026-03-14', '包含所有组件状态/断点/动效/深色模式'),
        ('每日站会标准化（15 分钟时间盒）', 'PM', '2026-03-12', '使用"昨日/今日/阻塞"三段式'),
        ('技术可行性评估表模板', 'BE', '2026-03-14', '需求评审前 1 天完成')
    ]
    for i, row in enumerate(w1_data, 1):
        for j, cell in enumerate(row):
            week1_table.rows[i].cells[j].text = cell
    
    doc.add_heading('6.2 第二周 (2026-03-19 至 2026-03-25)', level=2)
    week2_table = doc.add_table(rows=6, cols=4)
    week2_table.style = 'Table Grid'
    for j, h in enumerate(w1_headers):
        week2_table.rows[0].cells[j].text = h
    w2_data = [
        ('引入"需求评审会"机制', 'PM', '2026-03-18', '开发/QA/PM 三方确认'),
        ('建立"技术决策记录"(ADR) 模板', 'Arch', '2026-03-20', '包含背景/选项/决策/后果'),
        ('部署飞书机器人自动通知', 'BE', '2026-03-22', '任务状态变更自动推送'),
        ('搭建 Mock 服务器', 'BE', '2026-03-25', 'FE 可并行开发'),
        ('建立"发布检查清单"', 'QA', '2026-03-20', '回归测试/性能基准/回滚方案')
    ]
    for i, row in enumerate(w2_data, 1):
        for j, cell in enumerate(row):
            week2_table.rows[i].cells[j].text = cell
    
    doc.add_page_break()
    
    # 关键成功指标
    doc.add_heading('七、关键成功指标 (KPI)', level=1)
    kpi_table = doc.add_table(rows=9, cols=4)
    kpi_table.style = 'Table Grid'
    kpi_headers = ['指标', '当前值', '目标值', '测量方式']
    for j, h in enumerate(kpi_headers):
        kpi_table.rows[0].cells[j].text = h
    kpi_data = [
        ('需求变更返工率', '60% (3/5 次)', '< 20%', '变更追踪表统计'),
        ('联调问题数', '7 处', '< 2 处', 'QA 测试报告'),
        ('API 响应时间', '0.0078s', '< 0.05s', '性能测试'),
        ('页面加载时间', '< 0.01s', '< 2s', '性能测试'),
        ('缓存命中率', '100%', '> 90%', '后端监控'),
        ('Bug 报告完整率', '-', '100%', 'Bug 模板检查'),
        ('站会时长', '25 分钟', '< 15 分钟', '计时器'),
        ('紧急问题响应时间', '4 小时', '< 30 分钟', '监控告警')
    ]
    for i, row in enumerate(kpi_data, 1):
        for j, cell in enumerate(row):
            kpi_table.rows[i].cells[j].text = cell
    
    doc.add_page_break()
    
    # 总结
    doc.add_heading('八、总结', level=1)
    summary = doc.add_paragraph()
    summary.add_run('核心发现：').bold = True
    summary.add_run('\n技术问题背后往往是沟通问题，沟通问题背后往往是流程问题。\n\n')
    summary.add_run('改进关键：').bold = True
    summary.add_run('\n不在于引入更多工具，而在于严格执行已有的最佳实践，并建立持续改进的反馈闭环。\n\n')
    summary.add_run('对老大的承诺：').bold = True
    summary.add_run('\n作为团队，我们会严格执行上述改进措施，并在每周站会汇报进展。如有偏差，请直接批评。\n\n')
    
    # 落款
    doc.add_paragraph('\n' + '='*50)
    doc.add_paragraph('文档创建时间：2026-03-11')
    doc.add_paragraph('创建人：OCA 胡小豆')
    doc.add_paragraph('分发范围：全员 + 老大')
    
    # 保存文档
    output_path = '/home/admin/.openclaw/workspace/docs/AI_News_Pulse_项目复盘报告汇总.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_reflection_document()
